#!/usr/bin/env python3
"""
Resume Formatter - Matches original resume formatting
- Body Font: Calibri 10.5pt
- Header Font: Helvetica Neue Bold 11pt
- Tight spacing throughout
"""

import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ ERROR: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)

# --- FONT SPECS ---
FONT_HEADER = "Helvetica Neue"
FONT_BODY = "Calibri"

# Point Sizes
SIZE_NAME = 11           # Header name/title only
SIZE_BODY = 10.5         # Everything else
SIZE_SECTION = 10.5      # Section headers

# Spacing (in points)
MARGINS = 0.5
SPACE_AFTER_HEADER = 6
SPACE_BEFORE_SECT = 8
SPACE_AFTER_SECT = 2
SPACE_BEFORE_JOB = 6
SPACE_AFTER_JOB = 0
SPACE_AFTER_BULLET = 0
BULLET_INDENT = 0.25


def set_font(run, font_name, size, bold=False):
    """Apply font with XML fallbacks"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)


def add_bottom_border(paragraph):
    """Adds underline border to section headers"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """Set paragraph spacing"""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_bullet_paragraph(doc, indent=BULLET_INDENT):
    """Create a bullet point paragraph without using List Bullet style"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(SPACE_AFTER_BULLET)
    p.paragraph_format.space_before = Pt(0)

    # Add bullet character
    run = p.add_run("•\t")
    set_font(run, FONT_BODY, SIZE_BODY)

    return p


def create_resume_docx(data, output_path):
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(MARGINS)
        section.bottom_margin = Inches(MARGINS)
        section.left_margin = Inches(MARGINS)
        section.right_margin = Inches(MARGINS)

    # === HEADER: Name | Title ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0)

    name = data.get("name", "YOUR NAME").upper()
    run = p.add_run(name)
    set_font(run, FONT_HEADER, SIZE_NAME, bold=True)

    if data.get("title"):
        run = p.add_run(" | ")
        set_font(run, FONT_HEADER, SIZE_NAME, bold=False)
        run = p.add_run(data["title"])
        set_font(run, FONT_HEADER, SIZE_NAME, bold=False)

    # === CONTACT INFO ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=SPACE_AFTER_HEADER)

    contact = data.get("contact", {})
    parts = []
    if contact.get("location"): parts.append(contact["location"])
    if contact.get("phone"): parts.append(contact["phone"])
    if contact.get("email"): parts.append(contact["email"])

    run = p.add_run(" | ".join(parts))
    set_font(run, FONT_BODY, SIZE_BODY)

    # === SUMMARY ===
    if data.get("summary"):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=SPACE_AFTER_SECT)
        run = p.add_run(data["summary"])
        set_font(run, FONT_BODY, SIZE_BODY)

    # === SECTION HEADER HELPER ===
    def add_section_header(text):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=SPACE_BEFORE_SECT, after=SPACE_AFTER_SECT)
        run = p.add_run(text.upper())
        set_font(run, FONT_HEADER, SIZE_SECTION, bold=True)
        add_bottom_border(p)

    # === TECHNICAL SKILLS ===
    if data.get("skills"):
        add_section_header("Technical Skills")

        for category, skills in data.get("skills", {}).items():
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0)

            # Category (Bold)
            run = p.add_run(f"{category}: ")
            set_font(run, FONT_BODY, SIZE_BODY, bold=True)

            # Skills list
            run = p.add_run(", ".join(skills))
            set_font(run, FONT_BODY, SIZE_BODY)

    # === PROFESSIONAL EXPERIENCE ===
    if data.get("experiences"):
        add_section_header("Professional Experience")

        for exp in data.get("experiences", []):
            # Main title line: Title | COMPANY    Date
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=SPACE_BEFORE_JOB, after=0)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

            # Job title first (bold)
            if exp.get('title'):
                run = p.add_run(exp['title'])
                set_font(run, FONT_HEADER, SIZE_BODY, bold=True)
                run = p.add_run(" | ")
                set_font(run, FONT_BODY, SIZE_BODY)

            # Company name
            run = p.add_run(exp['company'].upper())
            set_font(run, FONT_BODY, SIZE_BODY, bold=True)

            # Tab + Date
            p.add_run("\t")
            run = p.add_run(exp['dates'])
            set_font(run, FONT_BODY, SIZE_BODY)

            # Previous title line (if exists)
            if exp.get('previousTitle'):
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=0)
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

                run = p.add_run(exp['previousTitle'])
                set_font(run, FONT_BODY, SIZE_BODY, bold=True)

                if exp.get('previousDates'):
                    p.add_run("\t")
                    run = p.add_run(exp['previousDates'])
                    set_font(run, FONT_BODY, SIZE_BODY)

            # Bullets
            for bullet in exp.get('bullets', []):
                p = add_bullet_paragraph(doc)

                text = bullet.get('text', '')
                title = bullet.get('title', '')

                if title:
                    run = p.add_run(f"{title}: ")
                    set_font(run, FONT_BODY, SIZE_BODY, bold=True)
                    run = p.add_run(text)
                    set_font(run, FONT_BODY, SIZE_BODY)
                elif ":" in text and text.index(":") < 50:
                    head, sep, tail = text.partition(":")
                    run = p.add_run(head + sep + " ")
                    set_font(run, FONT_BODY, SIZE_BODY, bold=True)
                    run = p.add_run(tail.strip())
                    set_font(run, FONT_BODY, SIZE_BODY)
                else:
                    run = p.add_run(text)
                    set_font(run, FONT_BODY, SIZE_BODY)

    # === EDUCATION ===
    if data.get("education"):
        add_section_header("Education & Certificates")

        for item in data.get("education", []):
            p = add_bullet_paragraph(doc)

            if isinstance(item, dict):
                text = item.get('school', '')
                if item.get('degree'): text += f", {item.get('degree')}"
                if item.get('year'): text += f" ({item.get('year')})"
            else:
                text = str(item)

            run = p.add_run(text)
            set_font(run, FONT_BODY, SIZE_BODY)

    doc.save(output_path)
    print(f"✅ Resume created: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 create_resume.py <data.json> <output.docx>")
        sys.exit(0)

    json_file = sys.argv[1]
    output_file = sys.argv[2]

    if not Path(json_file).exists():
        print(f"❌ Error: JSON file not found: {json_file}")
        sys.exit(1)

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        create_resume_docx(data, output_file)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
